# ============================================================
# parser/docx_parser.py — DOCX Resume Text Extraction
# ============================================================
# Uses python-docx to extract paragraphs and table cells.
# ============================================================

from __future__ import annotations

from pathlib import Path

from docx import Document

from utils.logger import logger


def parse_docx(file_path: str | Path) -> str:
    """
    Parse a DOCX resume and return extracted text.
    Extracts paragraphs and table cell text.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    logger.event("docx_parse_start", file=str(path))

    doc = Document(str(path))
    text_parts: list[str] = []

    # --- Paragraphs ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # --- Tables ---
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                text_parts.append(" | ".join(row_texts))

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        raise ValueError(f"No text found in DOCX: {path}")

    logger.event("docx_parse_done", file=str(path), chars=len(full_text))
    return full_text
